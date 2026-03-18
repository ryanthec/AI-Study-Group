from sqlalchemy.orm import Session, defer
from typing import List, Optional
import os
import io
import PyPDF2
import docx
from pathlib import Path
from datetime import datetime

from ..models.document_embedding import Document, DocumentChunk
from .embedding_service import embedding_service

class DocumentService:
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 2000, overlap: int = 400) -> List[str]:
        """
        Split text into overlapping chunks
        
        Args:
            text: Text to split
            chunk_size: Maximum characters per chunk
            overlap: Number of overlapping characters between chunks
        """
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            
            # Try to break at sentence boundary if possible
            if end < len(text):
                last_period = chunk.rfind('.')
                last_newline = chunk.rfind('\n')
                break_point = max(last_period, last_newline)
                
                if break_point > chunk_size * 0.5:  # Only break if we're past halfway
                    end = start + break_point + 1
                    chunk = text[start:end]
            
            chunks.append(chunk.strip())
            start = end - overlap
        
        return chunks
    
    @staticmethod
    def extract_text_from_bytes(filename: str, file_data: bytes) -> str:
        """
        Extract text content from file bytes without saving to DB.
        """
        content = ""
        filename = filename.lower()
        
        try:
            if filename.endswith('.pdf'):
                pdf_file = io.BytesIO(file_data)
                reader = PyPDF2.PdfReader(pdf_file)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        content += text + "\n"
                        
            elif filename.endswith('.docx'):
                docx_file = io.BytesIO(file_data)
                doc = docx.Document(docx_file)
                for para in doc.paragraphs:
                    content += para.text + "\n"
                    
            elif filename.endswith('.txt') or filename.endswith('.md'):
                content = file_data.decode('utf-8', errors='ignore')
                
        except Exception as e:
            print(f"Error extracting text from {filename}: {str(e)}")
            return f"[Error extracting text: {str(e)}]"
            
        # SANITIZE: Remove PostgreSQL-breaking NUL bytes
        return content.replace('\x00', '').replace('\0', '')
    
    @staticmethod
    def save_document_initial(
        db: Session, group_id: int, uploader_id: int, filename: str, 
        file_type: str, file_bytes: bytes, file_size: int
    ) -> Document:
        """Instantly save the raw file and metadata to the database."""

        if filename.lower().endswith('.docx') or 'wordprocessingml' in file_type.lower():
            # 1. Extract the text using your existing helper function
            extracted_text = DocumentService.extract_text_from_bytes(filename, file_bytes)
            
            # 2. Convert the clean text back into a standard byte stream
            file_bytes = extracted_text.encode('utf-8')
            
            # 3. Update the metadata so it is universally recognized as a text file
            filename = filename.rsplit('.', 1)[0] + '.txt'
            file_type = 'text/plain'
            file_size = len(file_bytes)

        document = Document(
            group_id=group_id,
            uploader_id=uploader_id,
            filename=filename,
            file_type=file_type,
            file_size=file_size,
            file_data=file_bytes,
            status="PENDING", # Set initial status
            created_at=datetime.utcnow()
        )
        db.add(document)
        db.commit()
        db.refresh(document)
        return document
    

    @staticmethod
    def process_document_background(document_id: int):
        """Heavy background task for extracting, chunking, and embedding."""
        from ..core.database import SessionLocal
        
        # Background tasks need a fresh database session
        db = SessionLocal()
        try:
            document = db.query(Document).filter(Document.id == document_id).first()
            if not document: return
            
            # 1. Extract text
            text_content = DocumentService.extract_text_from_bytes(document.filename, document.file_data)
            
            # 2. Chunk and Embed
            chunks = DocumentService.chunk_text(text_content)
            embeddings = embedding_service.embed_batch(chunks)
            
            # 3. Save Chunks
            for idx, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
                chunk = DocumentChunk(
                    document_id=document.id, chunk_index=idx, 
                    content=chunk_text, embedding=embedding, 
                    created_at=datetime.utcnow()
                )
                db.add(chunk)
            
            # 4. Mark as completed
            document.status = "COMPLETED"
            db.commit()
            
            # Optional: If you want to push a WebSocket notification to the frontend
            # from ..core.websocket_manager import manager
            # asyncio.run(manager.broadcast_to_group(document.group_id, {
            #     "type": "document_ready", "filename": document.filename
            # }))
            
        except Exception as e:
            print(f"Background processing failed for {document_id}: {e}")
            # CRITICAL: Rollback the poisoned transaction before trying to write again
            db.rollback() 
            
            # Now it's safe to update the status
            try:
                if document:
                    document.status = "ERROR"
                    db.commit()
            except Exception as inner_e:
                print(f"Failed to update error status: {inner_e}")
        finally:
            db.close()


    @staticmethod
    def similarity_search(
        db: Session,
        query_text: str,
        group_id: Optional[int] = None,
        limit: int = 5,
        similarity_threshold: float = 0.7
    ) -> List[DocumentChunk]:
        """
        Find most relevant document chunks for a query
        
        Args:
            db: Database session
            query_text: User's query
            group_id: Optional group ID to filter results
            limit: Maximum number of results
            similarity_threshold: Minimum similarity score (0-1)
        """
        # Generate embedding for query
        query_embedding = embedding_service.embed_text(query_text)
        
        # Apply threshold filter at database level
        results = db.query(DocumentChunk).join(Document).filter(
            Document.group_id == group_id,
            DocumentChunk.embedding.cosine_distance(query_embedding) <= (1 - similarity_threshold)
        ).order_by(
            DocumentChunk.embedding.cosine_distance(query_embedding)
        ).limit(limit).all()
        
        return results
    
    @staticmethod
    def delete_document(db: Session, document_id: int, group_id: int) -> bool:
        """Delete document and all its chunks"""
        document = db.query(Document).filter(
            Document.id == document_id,
            Document.group_id == group_id
        ).first()
        
        if document:
            # Cascade will delete chunks automatically
            db.delete(document)
            db.commit()
            return True
        
        return False


    @staticmethod
    def get_group_documents(db: Session, group_id: int, only_completed: bool = False):
        """Fetch documents. Agents should set only_completed=True to avoid reading pending files."""
        query = db.query(Document).options(defer(Document.file_data)).filter(
            Document.group_id == group_id
        )
        
        if only_completed:
            query = query.filter(Document.status == "COMPLETED")
            
        return query.all()

    @staticmethod
    def get_document_by_id(db: Session, document_id: int, group_id: int) -> Optional[Document]:
        """Get a specific document"""
        return db.query(Document).filter(
            Document.id == document_id,
            Document.group_id == group_id
        ).first()
    
    @staticmethod
    def get_document_with_bytes(db: Session, document_id: int, group_id: int) -> Optional[Document]:
        """Fetch a document including its heavy file_data binary."""
        return db.query(Document).filter(
            Document.id == document_id,
            Document.group_id == group_id
        ).first() # By default, SQLAlchemy won't defer unless explicitly told to
    